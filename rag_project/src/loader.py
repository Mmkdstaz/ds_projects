import io
import logging
import warnings
from pathlib import Path

warnings.filterwarnings("ignore")
logging.getLogger("transformers").setLevel(logging.ERROR)

from src.config import TESSERACT_CMD

logger = logging.getLogger(__name__)


def _has_text_layer(file_path: str, sample_pages: int = 5) -> bool:
    import fitz
    doc = fitz.open(file_path)
    total = min(sample_pages, len(doc))
    chars = sum(len(doc[i].get_text().strip()) for i in range(total))
    doc.close()
    return chars > total * 100


def load_pdf(file_path: str) -> list[dict]:
    import fitz
    import pytesseract
    from PIL import Image

    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

    doc = fitz.open(file_path)
    total = len(doc)
    use_ocr = not _has_text_layer(file_path)
    logger.info("%s — страниц: %d, OCR: %s", Path(file_path).name, total, use_ocr)

    pages = []
    for page_num in range(total):
        page = doc[page_num]
        text = page.get_text().strip()

        if use_ocr or len(text) < 20:
            pix = page.get_pixmap(dpi=300)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img, lang="rus+eng").strip()

        if text:
            pages.append({"text": text, "page": page_num + 1})

    doc.close()
    logger.info("Загружено страниц: %d/%d", len(pages), total)
    return pages


def load_docx(file_path: str) -> list[dict]:
    from docx import Document
    doc = Document(file_path)
    pages = []
    buffer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            buffer.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                buffer.append(" | ".join(cells))

    if buffer:
        pages.append({"text": "\n".join(buffer), "page": 1})

    return pages


def load_txt(file_path: str) -> list[dict]:
    try:
        text = Path(file_path).read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = Path(file_path).read_text(encoding="cp1251")
    return [{"text": text, "page": 1}]


def load_document(file_path: str) -> list[dict]:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return load_docx(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext}")